from docx import Document
import re

document = None
doc_name = None


def start_print(name):
    global document, doc_name
    assert re.match('^[\w-]{3,}$', name)
    document = Document()
    doc_name = name
    document.add_heading(f'Document for "{doc_name}"', 0)
    document.add_paragraph('A plain paragraph')


def end_print():
    assert document is not None
    document.save(f'MICS_{doc_name}.docx')


